import sys
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("python-docx not installed. Trying to install...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"])
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def md_to_docx(md_file, docx_file):
    doc = Document()
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    in_table = False
    
    for line in lines:
        line_clean = line.strip()
        
        # Code blocks
        if line_clean.startswith("```"):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line.rstrip())
            if p.runs:
                p.runs[0].font.name = 'Consolas'
                p.runs[0].font.size = Pt(9)
            continue
            
        # Headers
        if line_clean.startswith("# "):
            doc.add_heading(line_clean[2:], level=1)
        elif line_clean.startswith("## "):
            doc.add_heading(line_clean[3:], level=2)
        elif line_clean.startswith("### "):
            doc.add_heading(line_clean[4:], level=3)
        # Lists
        elif line_clean.startswith("- ") or line_clean.startswith("* "):
            doc.add_paragraph(line_clean[2:], style='List Bullet')
        # Empty lines
        elif not line_clean:
            # ignore empty lines if we already have space, or add simple paragraph
            pass
        # Normal text with basic bold
        else:
            p = doc.add_paragraph()
            parts = line_clean.split("**")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0:
                    run.bold = True
                    
    doc.save(docx_file)

if __name__ == "__main__":
    print("Converting...")
    md_to_docx("doc.md", "doc.docx")
    print("Done!")
